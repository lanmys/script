import fileinput
import  os
from docx import Document
from docx.shared import Pt            # 磅数



def files_name(file_dir):
    for root, dirs, files in os.walk(file_dir):
        print(files)       #当前路径下所有非目录子文件
    return files


def set_docx(file_name,file_cont):
    document = Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal'].font.size = Pt(5)
    text = document.add_paragraph()
    for i in file_cont:
        text.add_run(i)
    document.save(r'C:\Users\Administrator\Desktop\数据库巡检\{}.docx'.format(file_name))
    print('{}.docx文件创建完成'.format(file_name))


# file_name = files_name(r"C:\Users\Administrator\Desktop\数据库巡检")
#file = open(r'C:\Users\Administrator\Desktop\数据库巡检\2021-01-11.log','r',encoding='utf-8')
#file_cont = []
#while True:
#    file_line = file.readline()
#    print(file_line)
#    if not file_line:
#        break
#    file_cont.append(file_line)
#set_docx('2021-01-11.log',file_cont)



if __name__ == '__main__':
    file_name = files_name(r"C:\Users\Administrator\Desktop\数据库巡检")
    for i in file_name:
        file = open(r'C:\Users\Administrator\Desktop\数据库巡检\{}'.format(i),'r',encoding='utf-8')
        file_cont = []
        while True:
            file_line = file.readline()
            if not file_line:
                file.close()
                break
            file_cont.append(file_line)
        set_docx(i, file_cont)

